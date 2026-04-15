"""Extract text from DOCX files using python-docx."""
from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

logger = logging.getLogger("data_acq.extraction")


def extract_text_from_docx(docx_path: str | Path) -> str | None:
    try:
        path = Path(docx_path)
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            logger.warning("No text extracted from DOCX: %s", docx_path)
            return None

        return "\n".join(paragraphs)
    except Exception as e:
        logger.error("DOCX extraction failed for %s: %s", docx_path, e)
        return None
